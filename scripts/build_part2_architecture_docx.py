#!/usr/bin/env python3
from __future__ import annotations

import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
ASSETS_DIR = DOCS_DIR / "assets"
OUT_DIR = ROOT / "output" / "doc"

SOURCE_MD = DOCS_DIR / "PART2_ARCHITECTURE_FRAMEWORK_DESIGN.md"
RENDERED_MD = DOCS_DIR / "PART2_ARCHITECTURE_FRAMEWORK_DESIGN_RENDERED.md"
OUT_DOCX = OUT_DIR / "PART2_ARCHITECTURE_FRAMEWORK_DESIGN.docx"


UNICODE_DASHES = {
    "\u2010": "-",  # hyphen
    "\u2011": "-",  # non-breaking hyphen
    "\u2012": "-",  # figure dash
    "\u2013": "-",  # en dash
    "\u2014": "-",  # em dash
    "\u2212": "-",  # minus
}


def _normalise_text(text: str) -> str:
    for bad, good in UNICODE_DASHES.items():
        text = text.replace(bad, good)
    return text


def _run(cmd: list[str]) -> None:
    subprocess.run(cmd, cwd=str(ROOT), check=True)


def _add_inline_runs(paragraph, text: str) -> None:
    """
    Minimal markdown inline formatting: **bold**, *italic*, `code`.
    """
    pattern = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            continue
        if part.startswith("*") and part.endswith("*") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            continue
        if part.startswith("`") and part.endswith("`") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            continue

        paragraph.add_run(part)


def _set_document_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def _markdown_to_docx(markdown: str) -> Document:
    doc = Document()
    _set_document_defaults(doc)

    lines = markdown.splitlines()
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = " ".join(s.strip() for s in paragraph_buffer).strip()
        paragraph_buffer = []
        if not text:
            return
        p = doc.add_paragraph()
        _add_inline_runs(p, text)

    image_re = re.compile(r"!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)")

    for raw_line in lines:
        line = raw_line.rstrip("\n")

        # Skip markdown horizontal rules
        if line.strip() == "---":
            flush_paragraph()
            continue

        # Headings
        if line.startswith("# "):
            flush_paragraph()
            doc.add_paragraph(_normalise_text(line[2:].strip()), style="Title")
            continue
        if line.startswith("## "):
            flush_paragraph()
            doc.add_paragraph(_normalise_text(line[3:].strip()), style="Heading 1")
            continue
        if line.startswith("### "):
            flush_paragraph()
            doc.add_paragraph(_normalise_text(line[4:].strip()), style="Heading 2")
            continue

        # Images (rendered mermaid figures)
        m = image_re.match(line.strip())
        if m:
            flush_paragraph()
            rel = m.group("path").strip()
            img_path = (DOCS_DIR / rel).resolve() if not rel.startswith("/") else Path(rel)
            if not img_path.exists():
                raise FileNotFoundError(f"Image not found: {img_path}")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(img_path), width=Inches(6.2))
            continue

        # Blank line => end paragraph
        if not line.strip():
            flush_paragraph()
            continue

        # Normal paragraph text
        paragraph_buffer.append(_normalise_text(line))

    flush_paragraph()
    return doc


def main() -> None:
    # Ensure rendered markdown + figure PNGs are up to date.
    # This also guarantees assets/part2_figure_*.png exist.
    _run(["python3", "scripts/build_part2_architecture_pdf.py"])

    OUT_DIR.mkdir(parents=True, exist_ok=True)

    md_text = RENDERED_MD.read_text(encoding="utf-8")
    md_text = _normalise_text(md_text)

    doc = _markdown_to_docx(md_text)
    doc.save(str(OUT_DOCX))


if __name__ == "__main__":
    main()

